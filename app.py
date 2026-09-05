import os
import io
import json
import re
import time
import shutil
import streamlit as st
from google import genai
from google.genai import types, errors
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import gdown

st.set_page_config(page_title="ממשק כתיבת תל\"א", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Alef:wght@400;700&display=swap');

    /* החלת גופן אלף, גודל ברור ומרווח שורות מהודק ואלגנטי */
    html, body, .stMarkdown, p, h1, h2, h3, h4, label, input, textarea, button, select, [class*="css"], details, summary {
        font-family: 'Alef', sans-serif !important;
        direction: rtl !important;
        text-align: right !important;
        font-size: 1.12rem !important;
        line-height: 1.35 !important;
    }

    p {
        margin-top: 0 !important;
        margin-bottom: 0.35rem !important;
    }

    h1, .main-title {
        font-size: 2.3rem !important;
        font-weight: 700 !important;
        line-height: 1.2 !important;
    }

    h2, h3, .stSubheader {
        font-size: 1.55rem !important;
        font-weight: 700 !important;
        line-height: 1.25 !important;
        margin-top: 0.5rem !important;
        margin-bottom: 0.5rem !important;
    }

    /* כותרת ראשית – טקסט שחור, נקי ומודגש ללא מאפייני כפתור כלל */
    a.title-link, a.title-link:hover, a.title-link:visited, a.title-link:active, a.title-link:focus {
        text-decoration: none !important;
        color: #111111 !important;
        background: transparent !important;
        border: none !important;
        outline: none !important;
        box-shadow: none !important;
        display: inline-block !important;
        cursor: pointer !important;
    }
    a.title-link h1, a.title-link .main-title {
        color: #111111 !important;
        font-weight: 700 !important;
        font-size: 2.3rem !important;
        line-height: 1.2 !important;
        margin: 0 !important;
        padding: 0 !important;
        cursor: pointer !important;
        background: transparent !important;
        border: none !important;
    }

    /* כותרות חלונות המטרות */
    [data-testid="stExpander"] details summary p {
        font-size: 1.18rem !important;
        font-weight: 700 !important;
        color: #1a1a1a !important;
        line-height: 1.3 !important;
    }

    /* דגל יחיד מימין לכותרת המטרה */
    [data-testid="stExpander"] details summary p::before {
        content: "";
        display: inline-block;
        width: 18px;
        height: 18px;
        margin-left: 8px;
        vertical-align: -2px;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%23111111'%3E%3Cpath d='M6 3c-.55 0-1 .45-1 1v16c0 .55.45 1 1 1s1-.45 1-1v-6h11.5c.4 0 .75-.24.9-.6.15-.37.07-.79-.2-1.07L16.5 10l2.7-3.33c.27-.28.35-.7.2-1.07-.15-.36-.5-.6-.9-.6H7V4c0-.55-.45-1-1-1z'/%3E%3C/svg%3E");
        background-size: contain;
        background-repeat: no-repeat;
    }

    /* הסתרת הודעת 'Press Enter to apply' */
    [data-testid="InputInstructions"],
    .stTextInput small,
    div[data-testid="InputInstructions"] {
        display: none !important;
    }

    /* מלבן חיווי ירוק בהיר יחיד, צר, ללא בולד */
    div[data-testid="stAlert"] {
        background-color: #e8f5e9 !important;
        border: none !important;
        border-radius: 8px !important;
        width: fit-content !important;
        min-width: unset !important;
        max-width: fit-content !important;
        padding: 5px 15px !important;
        margin: 6px 0 !important;
        box-shadow: none !important;
    }
    div[data-testid="stAlert"] > div {
        background-color: transparent !important;
        border: none !important;
        padding: 0 !important;
        margin: 0 !important;
    }
    div[data-testid="stAlert"] p,
    div[data-testid="stAlert"] span,
    div[data-testid="stAlert"] * {
        color: #1b5e20 !important;
        font-weight: 400 !important;
        font-size: 1.05rem !important;
        line-height: 1.3 !important;
        background: transparent !important;
    }

    /* תיקון האייקונים של המערכת */
    [data-testid="stIcon"],
    [data-testid="stExpanderToggleIcon"],
    [class*="material-symbols"],
    [class*="material-icons"],
    span[class*="Icon"] {
        font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
        direction: ltr !important;
        text-align: left !important;
    }

    /* הסתרת סרגל ההגדרות בצד */
    [data-testid="stSidebar"] {
        display: none;
    }

    /* כפתור סטטוס דרייב תכלת */
    div[data-testid="stPopover"] {
        position: fixed !important;
        top: 12px !important;
        left: 20px !important;
        right: auto !important;
        width: auto !important;
        z-index: 999999 !important;
    }
    div[data-testid="stPopover"] > button {
        width: auto !important;
        min-width: unset !important;
        background-color: #f0f9ff !important;
        color: #0288D1 !important;
        border: 1px solid #bae6fd !important;
        border-radius: 20px !important;
        padding: 5px 15px !important;
        font-size: 0.96rem !important;
        font-weight: bold !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05) !important;
    }
    div[data-testid="stPopover"] > button:hover {
        background-color: #e0f2fe !important;
        border-color: #0288D1 !important;
    }
    div[data-testid="stPopoverBody"] {
        direction: rtl !important;
        text-align: right !important;
    }

    /* עיצוב כפתורים כללי בתכלת */
    .stButton>button, [data-testid="stFormSubmitButton"]>button {
        width: 100%;
        border-radius: 8px;
        font-weight: bold;
        font-family: 'Alef', sans-serif !important;
        direction: rtl !important;
        text-align: center !important;
        font-size: 1.02rem !important;
        padding: 6px 8px !important;
        white-space: nowrap !important;
        background-color: #f0f9ff !important;
        border: 1px solid #7dd3fc !important;
        color: #0369a1 !important;
    }
    .stButton>button:hover, [data-testid="stFormSubmitButton"]>button:hover {
        background-color: #bae6fd !important;
        border-color: #0288D1 !important;
        color: #0c4a6e !important;
    }

    /* כפתור ראשי תכלת מלא */
    .main-btn>button {
        background-color: #0288D1 !important;
        color: white !important;
        border: none !important;
        font-size: 1.2rem !important;
        padding: 10px !important;
    }
    .main-btn>button:hover {
        background-color: #0277BD !important;
        color: white !important;
    }

    /* מניעת רווחים מתגיות מרקרים */
    .marker-regen, .marker-edit, .marker-del, .marker-plus, .marker-download {
        display: none !important;
    }
    div[data-testid="element-container"]:has(.marker-regen),
    div[data-testid="element-container"]:has(.marker-edit),
    div[data-testid="element-container"]:has(.marker-del),
    div[data-testid="element-container"]:has(.marker-plus),
    div[data-testid="element-container"]:has(.marker-download),
    div.stElementContainer:has(.marker-regen),
    div.stElementContainer:has(.marker-edit),
    div.stElementContainer:has(.marker-del),
    div.stElementContainer:has(.marker-plus),
    div.stElementContainer:has(.marker-download) {
        margin: 0 !important;
        padding: 0 !important;
        height: 0 !important;
        min-height: 0 !important;
    }

    /* שילוב האייקונים בכפתורים */

    /* חץ עגול - נסח מחדש */
    div[data-testid="element-container"]:has(.marker-regen) + div[data-testid="element-container"] button p::before,
    div.stElementContainer:has(.marker-regen) + div.stElementContainer button p::before {
        content: "";
        display: inline-block;
        width: 16px;
        height: 16px;
        margin-left: 5px;
        vertical-align: -2px;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%230369a1'%3E%3Cpath d='M12 4V1L8 5l4 4V6c3.31 0 6 2.69 6 6 0 1.01-.25 1.97-.7 2.8l1.46 1.46C19.54 15.03 20 13.57 20 12c0-4.42-3.58-8-8-8zm0 14c-3.31 0-6-2.69-6-6 0-1.01.25-1.97.7-2.8L5.24 7.74C4.46 8.97 4 10.43 4 12c0 4.42 3.58 8 8 8v3l4-4-4-4v3z'/%3E%3C/svg%3E");
        background-size: contain;
        background-repeat: no-repeat;
    }

    /* עפרון - ערוך לפי תיאור */
    div[data-testid="element-container"]:has(.marker-edit) + div[data-testid="element-container"] button p::before,
    div.stElementContainer:has(.marker-edit) + div.stElementContainer button p::before {
        content: "";
        display: inline-block;
        width: 16px;
        height: 16px;
        margin-left: 5px;
        vertical-align: -2px;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%230369a1'%3E%3Cpath d='M3 17.25V21h3.75L17.81 9.94l-3.75-3.75L3 17.25zM20.71 7.04c.39-.39.39-1.02 0-1.41l-2.34-2.34c-.39-.39-1.02-.39-1.41 0l-1.83 1.83 3.75 3.75 1.83-1.83z'/%3E%3C/svg%3E");
        background-size: contain;
        background-repeat: no-repeat;
    }

    /* פח - מחק מטרה / מחק יעד */
    div[data-testid="element-container"]:has(.marker-del) + div[data-testid="element-container"] button p::before,
    div.stElementContainer:has(.marker-del) + div.stElementContainer button p::before {
        content: "";
        display: inline-block;
        width: 16px;
        height: 16px;
        margin-left: 5px;
        vertical-align: -2px;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%230369a1'%3E%3Cpath d='M6 19c0 1.1.9 2 2 2h8c1.1 0 2-.9 2-2V7H6v12zM19 4h-3.5l-1-1h-5l-1 1H5v2h14V4z'/%3E%3C/svg%3E");
        background-size: contain;
        background-repeat: no-repeat;
    }

    /* פלוס - הוסף יעד */
    div[data-testid="element-container"]:has(.marker-plus) + div[data-testid="element-container"] button p::before,
    div.stElementContainer:has(.marker-plus) + div.stElementContainer button p::before {
        content: "";
        display: inline-block;
        width: 17px;
        height: 17px;
        margin-left: 6px;
        vertical-align: -2px;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%230369a1'%3E%3Cpath d='M19 13h-6v6h-2v-6H5v-2h6V5h2v6h6v2z'/%3E%3C/svg%3E");
        background-size: contain;
        background-repeat: no-repeat;
    }

    /* כפתור הורדת Word - תכלת מלא */
    .download-btn>button {
        background-color: #0288D1 !important;
        color: #ffffff !important;
        border: none !important;
        font-size: 1.2rem !important;
        font-weight: 700 !important;
        padding: 10px !important;
    }
    .download-btn>button:hover {
        background-color: #0277BD !important;
        color: #ffffff !important;
    }
    div[data-testid="element-container"]:has(.marker-download) + div[data-testid="element-container"] button p::before,
    div.stElementContainer:has(.marker-download) + div.stElementContainer button p::before {
        content: "";
        display: inline-block;
        width: 20px;
        height: 20px;
        margin-left: 8px;
        vertical-align: -3px;
        background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='%23ffffff'%3E%3Cpath d='M19 9h-4V3H9v6H5l7 7 7-7zM5 18v2h14v-2H5z'/%3E%3C/svg%3E");
        background-size: contain;
        background-repeat: no-repeat;
    }

    /* תיקון אזור העלאת קבצים */
    [data-testid="stFileUploader"] section {
        direction: ltr !important;
        text-align: left !important;
    }
    [data-testid="stFileUploader"] section button {
        direction: ltr !important;
    }

    /* מתיחה מדויקת של חלון דרכי ההוראה */
    div.teach-box div[data-baseweb="textarea"],
    div.teach-box textarea,
    textarea[aria-label*="דרכי הוראה"] {
        height: 1000px !important;
        min-height: 1000px !important;
        max-height: 1000px !important;
        padding-top: 350px !important;
        padding-bottom: 40px !important;
        line-height: 2 !important;
        box-sizing: border-box !important;
        font-size: 1.1rem !important;
    }
</style>
""", unsafe_allow_html=True)

DRIVE_FOLDER_URL = "https://drive.google.com/drive/folders/1IHat-atuDDzFIfsmKq24aF7WkDjhbx0S?usp=drive_link"
LOCAL_DRIVE_FOLDER = "drive_examples"

@st.cache_data(show_spinner=False)
def sync_and_load_drive_examples():
    if os.path.exists(LOCAL_DRIVE_FOLDER):
        try:
            shutil.rmtree(LOCAL_DRIVE_FOLDER)
        except Exception:
            pass
    os.makedirs(LOCAL_DRIVE_FOLDER, exist_ok=True)
    
    try:
        gdown.download_folder(DRIVE_FOLDER_URL, output=LOCAL_DRIVE_FOLDER, quiet=True, use_cookies=False)
    except Exception:
        pass

    combined_examples = []
    if os.path.exists(LOCAL_DRIVE_FOLDER):
        for root, _, files in os.walk(LOCAL_DRIVE_FOLDER):
            for f in sorted(files):
                if f.endswith('.docx') and not f.startswith('~$'):
                    try:
                        doc = Document(os.path.join(root, f))
                        text_parts = []
                        for element in doc.element.body:
                            if element.tag.endswith('p'):
                                p = Paragraph(element, doc)
                                if p.text.strip():
                                    text_parts.append(p.text.strip())
                            elif element.tag.endswith('tbl'):
                                t = Table(element, doc)
                                for row in t.rows:
                                    row_text = " | ".join([c.text.replace("\n", " ").strip() for c in row.cells if c.text.strip()])
                                    if row_text:
                                        text_parts.append(row_text)
                                        
                        combined_examples.append(f"=== מאגר דוגמאות מתוך ({f}) ===\n" + "\n".join(text_parts))
                    except Exception:
                        pass
    return "\n\n".join(combined_examples)

def get_filtered_examples(full_context, target_class):
    """מסנן ומבודד מתוך מאגר הדוגמאות רק את הדוגמאות השייכות לכיתה המבוקשת (צעירים או בוגרים)"""
    if not full_context:
        return ""
    chunks = re.split(r'(?=(?:^|\n)\s*דוגמ[הא]\s*\d+)', full_context)
    filtered = []
    for chunk in chunks:
        lines = chunk.strip().splitlines()
        if not lines:
            continue
        first_line = lines[0]
        if "צעירים" in first_line or "בוגרים" in first_line:
            if target_class in first_line:
                filtered.append(chunk.strip())
        else:
            if not chunk.strip().startswith("=== מאגר דוגמאות"):
                filtered.append(chunk.strip())
    
    if filtered:
        return "\n\n".join(filtered)
    return full_context

def safe_parse_json(text_content):
    t = text_content.strip()
    if t.startswith("```"):
        lines = t.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        t = "\n".join(lines).strip()
    return json.loads(t)

def call_gemini_with_retry(client, contents, config=None, max_retries=3, initial_delay=2.0):
    delay = initial_delay
    last_err = None
    for attempt in range(max_retries):
        try:
            return client.models.generate_content(
                model='gemini-3.6-flash',
                contents=contents,
                config=config
            )
        except Exception as e:
            last_err = e
            is_server_err = isinstance(e, (errors.ServerError, errors.APIError))
            err_msg = str(e).upper()
            is_retryable = is_server_err or any(term in err_msg for term in ["503", "500", "502", "504", "429", "UNAVAILABLE", "RESOURCE_EXHAUSTED", "OVERLOADED"])
            
            if is_retryable and attempt < max_retries - 1:
                time.sleep(delay)
                delay *= 2
                continue
            raise e
    raise last_err

with st.spinner("מסנכרן דוגמאות מתיקיית הדרייב..."):
    examples_context = sync_and_load_drive_examples()

with st.popover("מאגר דרייב"):
    st.markdown("**סטטוס חיבור ל-Google Drive:**")
    if examples_context:
        num_docs = examples_context.count("=== מאגר דוגמאות מתוך")
        found_headers = re.findall(r'(?:^|\n|\b)דוגמ[הא]\s*(?:מס\'|מספר)?\s*(?:\d+|[א-ת](?:[\'״"][א-ת])?)\b', examples_context)
        num_examples = len(found_headers) if found_headers else num_docs
        num_young = len(re.findall(r'דוגמ[הא].*?צעירים', examples_context))
        num_old = len(re.findall(r'דוגמ[הא].*?בוגרים', examples_context))
        status_msg = f"מאגר הדוגמאות מחובר ומסונכרן!\n\nנטענו {num_docs} קבצים המכילים {num_examples} דוגמאות ללמידת המודל."
        if num_young or num_old:
            status_msg += f"\n\n(חלוקה: {num_young} דוגמאות צעירים, {num_old} דוגמאות בוגרים)"
        st.success(status_msg)
    else:
        st.warning("לא אותרו קבצים בתיקייה (וודאי שהשיתוף פתוח לצפייה לכולם).")

    if st.button("רענן מאגר דרייב", key="refresh_drive_btn"):
        st.cache_data.clear()
        st.rerun()

# כותרת ראשית – טקסט שחור, בולד וגדול (לחיצה עליה מרעננת ומאפסת למסך הבית)
st.markdown('<a href="./" target="_self" class="title-link"><h1 class="main-title">ממשק חכם לניסוח תל"א</h1></a>', unsafe_allow_html=True)
st.caption("מערכת לגזירת מטרות על ויעדים אופרטיביים על בסיס תיאור הילד/ה")

with st.sidebar:
    st.header("הגדרות מערכת")
    default_key = st.secrets.get("GEMINI_API_KEY", "")
    api_key = st.text_input(
        "מפתח API:", 
        value=default_key, 
        type="password"
    )

if 'goals_list' not in st.session_state:
    st.session_state['goals_list'] = []
if 'raw_uploaded_table' not in st.session_state:
    st.session_state['raw_uploaded_table'] = []

# טופס ראשי עם בחירת מגדר, כיתה ושם
col1, col2 = st.columns([1.2, 2.8])
with col1:
    gender = st.radio("התאמה מגדרית:", ["ילדה (נקבה)", "ילד (זכר)"])
    class_group = st.radio("כיתה:", ["צעירים", "בוגרים"])
with col2:
    student_name = st.text_input("שם הילד/ה:", value="", placeholder="הזיני את שם הילד/ה")

st.markdown("---")
st.subheader("1. תיאור רמת התפקוד בתחומים הרלוונטיים")
st.markdown("**הזנת מוקדי כוח ומוקדים לחיזוק**")

uploaded_file = st.file_uploader(
    "העלי קובץ Word של מוקדי כוח ומוקדים לחיזוק (יש לוודא שהקובץ סגור במחשב):", 
    type=["docx"]
)

input_text = ""
if uploaded_file is not None:
    try:
        doc_uploaded = Document(uploaded_file)
        extracted = [p.text for p in doc_uploaded.paragraphs if p.text.strip()]
        parsed_table_rows = []
        for table in doc_uploaded.tables:
            for row in table.rows:
                cells_text = [c.text.replace("\n", " ").strip() for c in row.cells]
                parsed_table_rows.append(cells_text)
                extracted.append(" | ".join(cells_text))
        st.session_state['raw_uploaded_table'] = parsed_table_rows
        input_text = "\n".join(extracted)
        st.success("הקובץ נטען בהצלחה.")
    except Exception as err:
        st.error("⚠️ לא ניתן לקרוא את הקובץ. אנא ודאי שמסמך ה-Word סגור במחשבך ולא נעול, ונסי להעלות שוב.")
else:
    input_text = st.text_area("או הדביקי כאן את הטקסט:", height=140, placeholder="הדביקי כאן מוקדי כוח וחיזוק...")

st.markdown('<div class="main-btn">', unsafe_allow_html=True)
if st.button("הפק מטרות ויעדים"):
    active_name = student_name.strip() if student_name.strip() else ("הילדה" if "נקבה" in gender else "הילד")
    if not input_text.strip():
        st.warning("אנא הזיני נתוני תפקוד.")
    else:
        with st.spinner(f"מנתח דפוסים ומנסח מטרות ויעדים תפקודיים המותאמים לכיתת {class_group}..."):
            try:
                env_pronoun = "לסביבתה" if "נקבה" in gender else "לסביבתו"
                
                # סינון הדוגמאות בהתאם לכיתה שנבחרה
                filtered_context = get_filtered_examples(examples_context, class_group)
                st.session_state['selected_class'] = class_group

                # הנחיות ייעודיות לצעירים מול בוגרים
                if class_group == "צעירים":
                    class_specific_instruction = """
- **התאמה בלעדית לכיתת צעירים (קריטי ומחייב):**
  * הילד/ה שייך/ת לכיתת צעירים. עליך להישען באופן בלעדי על סגנון הדוגמאות של הצעירים בלבד.
  * **איסור מוחלט על מטרות או יעדים של הכנה לכיתה א':** אין לנסח מטרות או יעדים העוסקים במודעות פונולוגית, זיהוי/כתיבת אותיות וספרות, קריאה, חשבון פורמלי, אנליזה וסינתזה או ישיבה ממושכת לדפי עבודה.
  * המטרות והיעדים יתמקדו בכישורי יסוד, תפקוד שגרתי יומיומי בגן, משחק הדדי ראשוני, הרחבת מיומנויות תקשורת והתפתחות מוטורית/שפתית תואמת גיל צעיר."""
                else:
                    class_specific_instruction = """
- **התאמה לכיתת בוגרים (קריטי ומחייב):**
  * הילד/ה שייך/ת לכיתת בוגרים. עליך להישען על סגנון הדוגמאות של הבוגרים (מיומנויות תפקוד מתקדמות, תפקודים ניהוליים, משחק חברתי עצמאי, פתרון קונפליקטים, שפה עשירה).
  * **שילוב מטרות הכנה לכיתה א' (מודעות פונולוגית, כתיבת אותיות/ספרות, קריאה ראשונית, התארגנות למטלה לימודית):** 
    יש לכלול מטרת הכנה לכיתה א' (או יעדים מותאמים לכך) **אך ורק אם מוקדי החיזוק של הילד/ה מצביעים בבירור על צורך וקושי בתחום זה** (כגון קשיים אורייניים, חשבון, תפקודים ניהוליים ללמידה או גרפומוטוריקה מתקדמת). אם מוקדי החיזוק של הילד/ה אינם מצריכים זאת – **אין להוסיף מטרת הכנה לכיתה א' באופן אוטומטי**, ויש להתמקד בתחומי התפקוד שבהם נדרש החיזוק."""

                system_prompt = f"""
אתה מומחה פדגוגי וקלינאי תקשורת בכיר לניסוח תכניות לימודים אישיות (תל\"א) בגני חינוך מיוחד.
עליך לנסח תל\"א מקצועית עבור '{active_name}' ({gender}, כיתת {class_group}) מתוך הישענות עמוקה על מאגר הדוגמאות הייעודי מתיקיית הדרייב.

### מאגר הדוגמאות המקצועיות ללמידה וחיקוי (מותאם לכיתת {class_group}):
{filtered_context}

---
### עקרונות מחייבים לניסוח מטרות ויעדים:
{class_specific_instruction}

1. **מטרת-על – רחבה, כוללת ותמציתית:**
   - שאיפה תפקודית רחבה (למשל: "{active_name} תביע כוונות תקשורתיות באמצעות משפטים...", "{active_name} תשתתף באופן מילולי במשימות חשיבה...", "{active_name} תנהל שיחה באופן הדדי...").
   - קצרה, ללא פירוט תנאים ספציפיים וללא חיבור שני תחומים ב-ו' החיבור.

2. **היעדים האופרטיביים – ספציפיים לתפקוד יחיד, חדים וללא סרבול (כלל קריטי):**
   - **כל יעד עוסק בתפקוד אחד ויחיד בלבד!**
   - **איסור מוחלט על העמסה וסרבול:** אין לדחוס מספר פעולות, תנאים ורמות תיווך באותו משפט (למשל, אין לכתוב: "{active_name} תשתמש במשפטים פשוטים להבעת צרכים, רצונות ורעיונות באופן המובן לסביבתה במהלך פעילויות השגרה והמשחק בגן, בתיווך מבוגר הולך ופוחת").
   - **יש לפרק לתפקודים בודדים ומדויקים, לדוגמה:**
     * יעד 1: "{active_name} תביע צרכים ורצונות באמצעות משפטים פשוטים"
     * יעד 2: "{active_name} תביע רעיונות במשחק עם מבוגר באמצעות שימוש במשפטים פשוטים"
     * יעד 3: "{active_name} תשתף בחוויה אישית קצרה באמצעות רצף משפטים פשוטים"
   - **בנושא מובנות דיבור והיגוי:** אם קיים קושי במובנות/היגוי, יש להקדיש לו **יעד נפרד וספציפי** (למשל: "{active_name} תהגה מילים דו-הברתיות באופן מובן בתוך שטף הדיבור"), ולא להעמיס את הביטוי "באופן המובן לסביבתה" על שאר יעדי השפה.

3. **התאמה מלאה למוקדי החיזוק של הילד/ה:**
   - גזור את המטרות והיעדים אך ורק מתוך תחומי הקושי שצוינו בטבלת התפקוד שהוזנה.

4. **דרכי הוראה, שיטות ואמצעים:** פירוט מעשי של אסטרטגיות מתוך שגרת הגן והטיפולים (שיחה בפת שחרית, משחקי קופסה, משחקי דמיון, מדרשי תמונה, מחברת שפה, טיפול פרטני/בזוגות, הדרכת הורים).

5. **כמות מחייבת:** בדיוק 3 מטרות-על. לכל מטרת-על בדיוק 3 יעדים ספציפיים.

---
### מבנה הפלט הנדרש (JSON בלבד):
[
  {{
    "goal_title": "{active_name} תשתתף / תפעל / תביע...",
    "domains": "תחום תפקוד",
    "objectives_list": [
      {{"text": "{active_name} [תפקוד יחיד ממוקד 1]...", "timeframe": "עד סוף השנה"}},
      {{"text": "{active_name} [תפקוד יחיד ממוקד 2]...", "timeframe": "עד סוף השנה"}},
      {{"text": "{active_name} [תפקוד יחיד ממוקד 3]...", "timeframe": "עד סוף השנה"}}
    ],
    "teaching_methods": "• אסטרטגיה 1\\n• אסטרטגיה 2\\n• אסטרטגיה 3"
  }}
]
"""
                client = genai.Client(api_key=api_key)
                response = call_gemini_with_retry(
                    client=client,
                    contents=f"נתוני התפקוד של {active_name} ({gender}, כיתת {class_group}):\n{input_text}",
                    config=types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        response_mime_type="application/json",
                        temperature=0.2
                    )
                )

                parsed_data = safe_parse_json(response.text)
                parsed_data = parsed_data[:3]
                for item in parsed_data:
                    item['ver'] = 0
                    if "objectives" in item and "objectives_list" not in item:
                        lines = [l.strip(" •\n\r") for l in item["objectives"].split("\n") if l.strip()]
                        item["objectives_list"] = [{"text": l, "timeframe": item.get("timeframe", "עד סוף השנה"), "ver": 0} for l in lines]

                    item_objs = item.get("objectives_list", [])[:3]
                    while len(item_objs) < 3:
                        item_objs.append({"text": f"{active_name} תשתתף בפעילות תפקודית מותאמת בהנחיית הצוות", "timeframe": "עד סוף השנה", "ver": 0})
                    for obj in item_objs:
                        obj['ver'] = 0
                    item['objectives_list'] = item_objs

                st.session_state['goals_list'] = parsed_data
                st.session_state['current_input_text'] = input_text
                st.session_state['just_generated'] = True
                st.rerun()
            except Exception as e:
                st.error(f"שגיאה בהפקה (עומס זמני בשרתים): {e}. אנא נסי שוב בעוד מספר שניות.")
st.markdown('</div>', unsafe_allow_html=True)

# חיווי הצלחה ירוק
if st.session_state.get('just_generated', False) and st.session_state.get('goals_list'):
    target_class_badge = st.session_state.get('selected_class', class_group)
    st.success(f"המטרות והיעדים הופקו בהצלחה בדגש תפקודי מותאם לכיתת {target_class_badge}!")

# ממשק עריכה אינטראקטיבי
if st.session_state['goals_list']:
    active_name = student_name.strip() if student_name.strip() else ("הילדה" if "נקבה" in gender else "הילד")
    current_class = st.session_state.get('selected_class', class_group)
    active_class_context = get_filtered_examples(examples_context, current_class)

    st.markdown("---")
    st.subheader("2. עריכה, דיוק והתאמת המטרות")

    for idx, goal in enumerate(st.session_state['goals_list']):
        g_ver = goal.get('ver', 0)
        current_title = goal.get('goal_title', '')

        with st.expander(f"מטרה {idx+1}: {current_title}", expanded=True, key=f"goal_expander_{idx}"):
            # 1. כותרת מטרה ותחומי תפקוד
            goal['goal_title'] = st.text_input(
                f"כותרת מטרה {idx+1}:", 
                value=current_title, 
                key=f"title_{idx}_{g_ver}"
            )

            goal['domains'] = st.text_input(
                f"תחומי תפקוד:", 
                value=goal.get('domains', ''), 
                key=f"dom_{idx}_{g_ver}"
            )

            # שורת כפתורי פעולה למטרה
            c_g_box, _ = st.columns([5, 5], gap="small")
            with c_g_box:
                c_g1, c_g2, c_g3 = st.columns([1.1, 1.35, 1.0], gap="small")
                with c_g1:
                    st.markdown('<span class="marker-regen"></span>', unsafe_allow_html=True)
                    if st.button("נסח מחדש", key=f"btn_regen_goal_{idx}"):
                        with st.spinner("מנסח חלופה תפקודית כוללת למטרה..."):
                            try:
                                client = genai.Client(api_key=api_key)
                                regen_prompt = f"""אתה מומחה לניסוח תל\"א בגני חינוך מיוחד.
הצע ניסוח חלופי, כללי ותפקודי (מוכוון השתתפות פעילה בשגרת הגן) למטרת-העל עבור {active_name} ({gender}, כיתת {current_class}).
התבסס באופן הדוק על הסגנון והשפה במאגר הדוגמאות של כיתת {current_class}:
{active_class_context}

הניסוח הנוכחי: '{goal['goal_title']}'
רקע נתוני תפקוד:
{st.session_state.get('current_input_text', '')}

דגשים קריטיים:
1. ניסוח תפקודי ברוח הדוגמאות של כיתת {current_class} (למשל: '{active_name} תשתתף...', '{active_name} תביע...', '{active_name} תיקח חלק...'). אסור לנסח 'תרכוש מיומנות'.
2. מטרת-על כללית וכוללת, קצרה ותמציתית, ללא תנאים ספציפיים בכותרת.
3. השתמש בשם המפורש '{active_name}'.
4. החזר אך ורק מחרוזת טקסט פשוטה של המטרה ללא מרכאות."""
                                res = call_gemini_with_retry(
                                    client=client,
                                    contents=regen_prompt,
                                    config=types.GenerateContentConfig(temperature=0.2)
                                )
                                goal['goal_title'] = res.text.strip().replace('"', '').replace("'", "")
                                goal['ver'] = g_ver + 1
                                st.rerun()
                            except Exception as e:
                                st.error(f"שגיאה בניסוח מחדש: {e}")

                with c_g2:
                    st.markdown('<span class="marker-edit"></span>', unsafe_allow_html=True)
                    if st.button("ערוך לפי תיאור", key=f"btn_toggle_edit_g_{idx}"):
                        st.session_state[f"show_edit_g_{idx}"] = not st.session_state.get(f"show_edit_g_{idx}", False)
                        st.rerun()

                with c_g3:
                    st.markdown('<span class="marker-del"></span>', unsafe_allow_html=True)
                    if st.button("מחק מטרה", key=f"del_{idx}"):
                        st.session_state['goals_list'].pop(idx)
                        st.rerun()

            # חלון עריכת מטרה שנפתח רק בלחיצה
            if st.session_state.get(f"show_edit_g_{idx}", False):
                c_form_g, _ = st.columns([5, 5], gap="small")
                with c_form_g:
                    with st.form(key=f"form_edit_g_{idx}", clear_on_submit=False, border=False):
                        f_col1, f_col2 = st.columns([4, 1], gap="small")
                        with f_col1:
                            prompt_g_val = st.text_input(
                                "תיאור לעריכת מטרה:", 
                                placeholder="למשל: התייחס למובנות הדיבור", 
                                key=f"edit_g_p_{idx}", 
                                label_visibility="collapsed"
                            )
                        with f_col2:
                            submit_edit_g = st.form_submit_button("שלח")

                        if submit_edit_g and prompt_g_val.strip():
                            with st.spinner("מעדכן ניסוח מטרה..."):
                                try:
                                    client = genai.Client(api_key=api_key)
                                    res = call_gemini_with_retry(
                                        client=client,
                                        contents=f"ערוך את מטרת-העל עבור {active_name} ({gender}, כיתת {current_class}): '{goal['goal_title']}' לפי ההנחיה: '{prompt_g_val}'. שמור על סגנון הדוגמאות של כיתת {current_class} מתיקיית הדרייב ועל ניסוח תפקודי כולל, קצר ובהיר. השתמש בשם המפורש '{active_name}'. החזר טקסט בלבד.",
                                        config=types.GenerateContentConfig(temperature=0.2)
                                    )
                                    goal['goal_title'] = res.text.strip().replace('"', '').replace("'", "")
                                    goal['ver'] = g_ver + 1
                                    st.session_state[f"show_edit_g_{idx}"] = False
                                    st.session_state[f"edit_g_p_{idx}"] = ""
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"שגיאה בעדכון מטרה: {e}")

            st.markdown("---")

            # מבנה טבלאי של היעדים ודרכי ההוראה
            t_col_left, t_col_right = st.columns([6, 4])
            with t_col_left:
                for o_idx, obj_item in enumerate(goal.get('objectives_list', [])):
                    o_ver = obj_item.get('ver', 0)
                    col_obj_text, col_obj_tf = st.columns([4.2, 1.8], gap="small")
                    with col_obj_text:
                        st.markdown(f"**יעד {o_idx+1}:**")
                        obj_item['text'] = st.text_area(
                            f"טקסט יעד {o_idx+1}", 
                            value=obj_item.get('text', ''), 
                            height=85, 
                            label_visibility="collapsed", 
                            key=f"obj_txt_{idx}_{o_idx}_{o_ver}"
                        )

                        # שורת כפתורי פעולה ליעד
                        c_b1, c_b2, c_b3 = st.columns([1.1, 1.35, 1.0], gap="small")
                        with c_b1:
                            st.markdown('<span class="marker-regen"></span>', unsafe_allow_html=True)
                            if st.button("נסח מחדש", key=f"btn_reg_obj_{idx}_{o_idx}"):
                                with st.spinner("מנסח יעד ספציפי וממוקד..."):
                                    try:
                                        client = genai.Client(api_key=api_key)
                                        regen_obj_prompt = f"""אתה מומחה לניסוח יעדים אופרטיביים בתל\"א לגני חינוך מיוחד.
הצע ניסוח חלופי ליעד אופרטיבי זה בלבד עבור {active_name} ({gender}, כיתת {current_class}) הנגזר ממטרת-העל '{goal['goal_title']}'.
התבסס באופן מלא על שפת הדוגמאות של כיתת {current_class} במאגר:
{active_class_context}

היעד הנוכחי: '{obj_item.get('text', '')}'
רקע נתוני תפקוד:
{st.session_state.get('current_input_text', '')}

דגשים מחייבים:
1. **תפקוד יחיד וספציפי בלבד:** על היעד להתמקד בפעולה מדויקת אחת (ללא העמסה וללא סרבול).
2. התאמה לרמת כיתת {current_class}.
3. ניסוח בהיר, קצר וישיר שפותח בשם המפורש '{active_name}'.
4. איסור על ניסוח 'תרכוש מיומנות' - השתמש בפועל של עשייה והשתתפות.
5. החזר משפט יחיד בלבד ללא מרכאות או תוספות."""
                                        res = call_gemini_with_retry(
                                            client=client,
                                            contents=regen_obj_prompt,
                                            config=types.GenerateContentConfig(temperature=0.2)
                                        )
                                        obj_item['text'] = res.text.strip().replace('"', '').replace("'", "")
                                        obj_item['ver'] = o_ver + 1
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"שגיאה בניסוח יעד: {e}")

                        with c_b2:
                            st.markdown('<span class="marker-edit"></span>', unsafe_allow_html=True)
                            if st.button("ערוך לפי תיאור", key=f"btn_toggle_pr_obj_{idx}_{o_idx}"):
                                st.session_state[f"show_pr_obj_{idx}_{o_idx}"] = not st.session_state.get(f"show_pr_obj_{idx}_{o_idx}", False)
                                st.rerun()

                        with c_b3:
                            st.markdown('<span class="marker-del"></span>', unsafe_allow_html=True)
                            if st.button("מחק יעד", key=f"del_obj_{idx}_{o_idx}"):
                                goal['objectives_list'].pop(o_idx)
                                st.rerun()

                        # חלון עריכת יעד שנפתח רק בלחיצה
                        if st.session_state.get(f"show_pr_obj_{idx}_{o_idx}", False):
                            with st.form(key=f"form_pr_obj_{idx}_{o_idx}", clear_on_submit=False, border=False):
                                f_col1, f_col2 = st.columns([4, 1], gap="small")
                                with f_col1:
                                    prompt_obj_val = st.text_input(
                                        "ערוך יעד לפי תיאור:", 
                                        placeholder="למשל: נסח בצורה פשוטה יותר", 
                                        key=f"pr_obj_{idx}_{o_idx}", 
                                        label_visibility="collapsed"
                                    )
                                with f_col2:
                                    submit_pr_obj = st.form_submit_button("שלח")

                                if submit_pr_obj and prompt_obj_val.strip():
                                    with st.spinner("מעדכן יעד..."):
                                        try:
                                            client = genai.Client(api_key=api_key)
                                            res = call_gemini_with_retry(
                                                client=client,
                                                contents=f"ערוך את היעד של {active_name} ({gender}, כיתת {current_class}): '{obj_item.get('text', '')}' לפי ההנחיה: '{prompt_obj_val}'. ודא שהיעד נגזר ממטרת-העל: '{goal['goal_title']}', עוסק בתפקוד יחיד ומוגדר בלבד, ללא סרבול, מותאם לכיתת {current_class}, ופותח בשם '{active_name}'. החזר משפט יחיד בלבד.",
                                                config=types.GenerateContentConfig(temperature=0.2)
                                            )
                                            obj_item['text'] = res.text.strip().replace('"', '').replace("'", "")
                                            obj_item['ver'] = o_ver + 1
                                            st.session_state[f"show_pr_obj_{idx}_{o_idx}"] = False
                                            st.session_state[f"pr_obj_{idx}_{o_idx}"] = ""
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"שגיאה בעדכון יעד: {e}")

                    with col_obj_tf:
                        st.markdown("**פרק זמן להשגה**")
                        obj_item['timeframe'] = st.text_input(
                            f"פרק זמן ליעד {o_idx+1}", 
                            value=obj_item.get('timeframe', 'עד סוף השנה'), 
                            key=f"obj_T_{idx}_{o_idx}_{o_ver}", 
                            label_visibility="collapsed"
                        )

                    st.markdown("---")

                # הוספת יעד חדש
                st.markdown("**הוספת יעד חדש:**")
                c_add_btn_box, _ = st.columns([5, 5], gap="small")
                with c_add_btn_box:
                    st.markdown('<span class="marker-plus"></span>', unsafe_allow_html=True)
                    if st.button("הוסף יעד למטרה זו", key=f"do_add_obj_auto_{idx}"):
                        with st.spinner("מנסח יעד תפקודי חדש וממוקד..."):
                            try:
                                client = genai.Client(api_key=api_key)
                                prompt_add = f"""אתה מומחה לניסוח תל\"א.
הוסף יעד אופרטיבי נוסף, ספציפי וממוקד בתפקוד יחיד, שנגזר ישירות ממטרת-העל: '{goal['goal_title']}' עבור {active_name} ({gender}, כיתת {current_class}).
התבסס על הסגנון והטרמינולוגיה בדוגמאות של כיתת {current_class}:
{active_class_context}

הקפד לפתוח בשם המפורש '{active_name}', לשמור על ניסוח קצר וממוקד בתפקוד יחיד ללא סרבול. החזר משפט יחיד בלבד ללא מרכאות."""
                                res = call_gemini_with_retry(
                                    client=client,
                                    contents=prompt_add,
                                    config=types.GenerateContentConfig(temperature=0.2)
                                )
                                goal.setdefault('objectives_list', []).append({
                                    "text": res.text.strip().replace('"', '').replace("'", ""), 
                                    "timeframe": "עד סוף השנה",
                                    "ver": 0
                                })
                                st.rerun()
                            except Exception as e:
                                st.error(f"שגיאה בהוספת יעד: {e}. אנא נסי שוב.")

                # שורת הוספת יעד על פי תיאור עם כפתור שלח
                with st.form(key=f"form_add_obj_custom_{idx}", clear_on_submit=False, border=False):
                    f_col1, f_col2 = st.columns([5, 1], gap="small")
                    with f_col1:
                        add_obj_val = st.text_input(
                            "הוסף יעד על פי תיאור:", 
                            placeholder="הוסף יעד על פי תיאור. למשל: הוסף יעד בסיסי יותר", 
                            key=f"add_obj_p_{idx}", 
                            label_visibility="collapsed"
                        )
                    with f_col2:
                        submit_add_obj = st.form_submit_button("שלח")

                    if submit_add_obj and add_obj_val.strip():
                        with st.spinner("מנסח יעד תפקודי חדש וממוקד..."):
                            try:
                                client = genai.Client(api_key=api_key)
                                prompt_add = f"""אתה מומחה לניסוח תל\"א.
הוסף יעד אופרטיבי נוסף, ספציפי וממוקד בתפקוד יחיד, שנגזר ישירות ממטרת-העל: '{goal['goal_title']}' עבור {active_name} ({gender}, כיתת {current_class}).
התבסס על הסגנון והטרמינולוגיה בדוגמאות של כיתת {current_class}:
{active_class_context}

דגש מיוחד ליעד: {add_obj_val}.
הקפד לפתוח בשם המפורש '{active_name}', לשמור על ניסוח קצר וממוקד בתפקוד יחיד ללא סרבול. החזר משפט יחיד בלבד ללא מרכאות."""
                                res = call_gemini_with_retry(
                                    client=client,
                                    contents=prompt_add,
                                    config=types.GenerateContentConfig(temperature=0.2)
                                )
                                goal.setdefault('objectives_list', []).append({
                                    "text": res.text.strip().replace('"', '').replace("'", ""), 
                                    "timeframe": "עד סוף השנה",
                                    "ver": 0
                                })
                                st.session_state[f"add_obj_p_{idx}"] = ""
                                st.rerun()
                            except Exception as e:
                                st.error(f"שגיאה בהוספת יעד: {e}. אנא נסי שוב.")

            with t_col_right:
                st.markdown("**דרכי הוראה, השיטות והאמצעים**")
                st.markdown('<div class="teach-box">', unsafe_allow_html=True)
                goal['teaching_methods'] = st.text_area(
                    "דרכי הוראה ואמצעים:", 
                    value=goal.get('teaching_methods', ''), 
                    height=1000, 
                    key=f"teach_{idx}_{g_ver}", 
                    label_visibility="collapsed"
                )
                st.markdown('</div>', unsafe_allow_html=True)

    # הוספת מטרה נוספת
    st.markdown("#### הוספת מטרה נוספת")
    with st.form(key="form_add_goal", clear_on_submit=False, border=False):
        col_new1, col_new2 = st.columns([3, 1], gap="small")
        with col_new1:
            custom_prompt = st.text_input(
                "איזו מטרה תרצי להוסיף?", 
                placeholder="למשל: הוסף מטרה בתחום המשחק", 
                key="new_goal_inp", 
                label_visibility="collapsed"
            )
        with col_new2:
            st.markdown('<span class="marker-plus"></span>', unsafe_allow_html=True)
            submit_add_goal = st.form_submit_button("הוסף מטרה זו")

        if submit_add_goal and custom_prompt.strip():
            with st.spinner(f"מנסח מטרה ויעדים תפקודיים על בסיס דוגמאות {current_class}..."):
                try:
                    client = genai.Client(api_key=api_key)
                    add_prompt = f"""
אתה מומחה לניסוח תל\"א בגני חינוך מיוחד.
נסח מטרה חדשה בתחום '{custom_prompt}' עבור {active_name} ({gender}, כיתת {current_class}) על בסיס מאגר הדוגמאות של כיתת {current_class}:
{active_class_context}

כללים:
1. מטרת-על כללית וכוללת בדגש על השתתפות ותפקוד יומיומי ברוח דוגמאות ה{current_class} עבור '{active_name}' (ללא 'תרכוש מיומנות').
2. בדיוק 3 יעדים אופרטיביים ממוקדים (כל יעד עוסק בתפקוד יחיד וברור, ללא סרבול והעמסה) שנגזרים ממנה ופותחים בשם '{active_name}'.
3. דרכי הוראה טיפוליות מעשיות.

החזר JSON יחיד בלבד במבנה:
{{
  "goal_title": "{active_name} תשתתף / תפעל...",
  "domains": "תחום תפקוד",
  "objectives_list": [
    {{"text": "{active_name} [תפקוד יחיד ממוקד 1]...", "timeframe": "עד סוף השנה"}},
    {{"text": "{active_name} [תפקוד יחיד ממוקד 2]...", "timeframe": "עד סוף השנה"}},
    {{"text": "{active_name} [תפקוד יחיד ממוקד 3]...", "timeframe": "עד סוף השנה"}}
  ],
  "teaching_methods": "• דרך הוראה 1\\n• דרך הוראה 2"
}}
"""
                    res = call_gemini_with_retry(
                        client=client,
                        contents=add_prompt,
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            temperature=0.2
                        )
                    )
                    new_item = safe_parse_json(res.text)
                    if isinstance(new_item, list) and len(new_item) > 0:
                        new_item = new_item[0]
                    new_item['ver'] = 0
                    for obj in new_item.get('objectives_list', []):
                        obj['ver'] = 0
                    st.session_state['goals_list'].append(new_item)
                    st.rerun()
                except Exception as e:
                    st.error(f"שגיאה בהוספת מטרה: {e}. אנא נסי שוב.")

    # ייצוא קובץ Word
    st.markdown("---")
    st.subheader("3. ייצוא המסמך הסופי")

    def apply_pPr_rtl(p, align="right", space_after_pt=3, line_spacing=1.15):
        val_align = "center" if (align == "center" or align == WD_ALIGN_PARAGRAPH.CENTER) else "right"
        sp_after = str(int(space_after_pt * 20)) if space_after_pt is not None else "60"
        ln_sp = str(int(line_spacing * 240)) if line_spacing is not None else "276"

        pPr_xml = (
            f'<w:pPr {nsdecls("w")}>'
            f'<w:bidi/>'
            f'<w:spacing w:after="{sp_after}" w:line="{ln_sp}" w:lineRule="auto"/>'
            f'<w:jc w:val="{val_align}"/>'
            f'</w:pPr>'
        )
        pPr_new = parse_xml(pPr_xml)
        if p._p.pPr is not None:
            p._p.remove(p._p.pPr)
        p._p.insert(0, pPr_new)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if val_align == "center" else WD_ALIGN_PARAGRAPH.RIGHT

    def apply_rPr_david_rtl(run, bold=None, size_pt=None, italic=False, underline=False):
        is_bold = bold if bold is not None else bool(run.bold)
        sz_pt = size_pt if size_pt is not None else (run.font.size.pt if run.font.size else 11.0)
        half_sz = str(int(sz_pt * 2))

        b_tag = "<w:b/><w:bCs/>" if is_bold else ""
        i_tag = "<w:i/><w:iCs/>" if italic else ""
        u_tag = '<w:u w:val="single"/>' if underline else ""

        rPr_xml = (
            f'<w:rPr {nsdecls("w")}>'
            f'<w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>'
            f'{b_tag}'
            f'{i_tag}'
            f'<w:sz w:val="{half_sz}"/>'
            f'<w:szCs w:val="{half_sz}"/>'
            f'{u_tag}'
            f'<w:rtl/>'
            f'<w:lang w:bidi="he-IL"/>'
            f'</w:rPr>'
        )
        rPr_new = parse_xml(rPr_xml)
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)
        run._r.insert(0, rPr_new)
        run.font.name = 'David'
        run.font.size = Pt(sz_pt)
        run.bold = is_bold

    def set_cell_clean_text(cell, text, bold=False, size_pt=11.0, align="right"):
        cell.text = ""
        lines = str(text).split("\n")
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            r = p.add_run(line)
            apply_rPr_david_rtl(r, bold=bold, size_pt=size_pt)
            apply_pPr_rtl(p, align=align, space_after_pt=2)

        tcPr = cell._tc.get_or_add_tcPr()
        for el in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(el)
        tcPr.append(parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>'))

    def build_template_docx():
        template_candidates = ["template.docx", "מבנה הקובץ.docx", "מבנה הקובץ_2.docx", "מבנה הקובץ_3.docx", "template_tala.docx"]
        existing_template = next((p for p in template_candidates if os.path.exists(p)), None)

        if existing_template:
            doc = Document(existing_template)

            for section in doc.sections:
                sectPr = section._sectPr
                for el in sectPr.findall(qn('w:bidi')):
                    sectPr.remove(el)
                sectPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))

            # 1. עדכון שם הילד
            child_label_name = active_name if active_name not in ["הילדה", "הילד", ""] else "________________"
            for p in doc.paragraphs:
                if "שם הילד" in p.text:
                    p.text = f"שם הילד/ה:  {child_label_name}       ת. לידה: _______________       ת.ז: ________________"
                    apply_pPr_rtl(p, align="right", space_after_pt=12)
                    for r in p.runs:
                        apply_rPr_david_rtl(r, bold=True, size_pt=12.0)
                    break

            # 2. שיבוץ מטרות ותחומי תפקוד
            for idx, g in enumerate(st.session_state['goals_list']):
                target_goal_num = f"מטרה {idx+1}:"
                for p_i, p in enumerate(doc.paragraphs):
                    if target_goal_num in p.text and "תחומי" not in p.text:
                        p.text = f"מטרה {idx+1}: {g.get('goal_title', '')}"
                        apply_pPr_rtl(p, align="right", space_after_pt=3)
                        for r in p.runs:
                            apply_rPr_david_rtl(r, bold=True, size_pt=13.0)

                        for next_p in doc.paragraphs[p_i+1:p_i+4]:
                            if "תחומי התפקוד" in next_p.text or "תחומי תפקוד" in next_p.text:
                                dom_val = g.get('domains', '').strip()
                                if not dom_val:
                                    dom_val = "קוגניטיבי / לימודי / התנהגותי רגשי / חברתי / חושי / מוטורי / תקשורתי-שפתי / כישורי חיים"
                                next_p.text = ""
                                r_pref = next_p.add_run("תחומי התפקוד אליהם מתייחסת : ")
                                apply_rPr_david_rtl(r_pref, bold=True, size_pt=11.5)
                                r_val = next_p.add_run(dom_val)
                                apply_rPr_david_rtl(r_val, bold=False, size_pt=11.5)
                                apply_pPr_rtl(next_p, align="right", space_after_pt=6)
                                break
                        break

            # 3. עדכון טבלת התפקוד (3 עמודות) - עמודה ימנית (תחום) בבולד
            func_tables = [t for t in doc.tables if len(t.columns) == 3]
            raw_rows = st.session_state.get('raw_uploaded_table', [])
            valid_uploaded_rows = [r for r in raw_rows if len(r) >= 3 and not all(c.strip() == "" for c in r)]

            if func_tables and len(valid_uploaded_rows) > 1:
                t_func = func_tables[0]
                for r_i, r_data in enumerate(valid_uploaded_rows[1:]):
                    t_idx = r_i + 1
                    if t_idx >= len(t_func.rows):
                        t_func.add_row()
                    row = t_func.rows[t_idx]
                    for c_i in range(min(3, len(r_data))):
                        set_cell_clean_text(row.cells[c_i], r_data[c_i].strip(), bold=(c_i == 0), size_pt=11.0, align="right")

            # 4. שיבוץ יעדים, פרק זמן ודרכי הוראה בטבלאות המטרות (5 עמודות)
            goal_tables = [t for t in doc.tables if len(t.columns) == 5]
            for idx, g in enumerate(st.session_state['goals_list']):
                if idx < len(goal_tables):
                    t_goal = goal_tables[idx]

                    header_cells = [c.text.strip() for c in t_goal.rows[0].cells]
                    obj_col, tf_col, teach_col = 0, 1, 2
                    for ci, h_txt in enumerate(header_cells):
                        if "יעד" in h_txt:
                            obj_col = ci
                        elif "פרק זמן" in h_txt:
                            tf_col = ci
                        elif "דרכי הוראה" in h_txt or "שיטות" in h_txt:
                            teach_col = ci

                    objs = g.get('objectives_list', [])
                    n_objs = max(len(objs), 1)
                    teach_txt = g.get('teaching_methods', '').strip()

                    while len(t_goal.rows) < n_objs + 1:
                        t_goal.add_row()

                    for o_i, obj_item in enumerate(objs):
                        row = t_goal.rows[o_i + 1]

                        raw_text = obj_item.get('text', '')
                        clean_obj = re.sub(r'^[\s•\-\*\d\.\)]+', '', raw_text).strip()
                        set_cell_clean_text(row.cells[obj_col], clean_obj, bold=False, size_pt=11.0, align="right")

                        tf_val = obj_item.get('timeframe', 'עד סוף השנה').strip()
                        set_cell_clean_text(row.cells[tf_col], tf_val, bold=False, size_pt=11.0, align="center")

                    if n_objs > 1:
                        c_top = t_goal.cell(1, teach_col)
                        c_bot = t_goal.cell(n_objs, teach_col)
                        c_merged = c_top.merge(c_bot)
                        set_cell_clean_text(c_merged, teach_txt, bold=False, size_pt=10.5, align="right")

                        other_cols = [c for c in range(5) if c not in (obj_col, tf_col, teach_col)]
                        for oc in other_cols:
                            t_goal.cell(1, oc).merge(t_goal.cell(n_objs, oc))
                    else:
                        set_cell_clean_text(t_goal.cell(1, teach_col), teach_txt, bold=False, size_pt=10.5, align="right")

            # 5. החלת פונט David, Bidi ויישור ימני מלא על כלל הפסקאות והטבלאות במסמך
            for p in doc.paragraphs:
                is_centered = (p.alignment == WD_ALIGN_PARAGRAPH.CENTER) or ("תוכנית לימודים" in p.text)
                target_align = "center" if is_centered else "right"
                for r in p.runs:
                    apply_rPr_david_rtl(r)
                apply_pPr_rtl(p, align=target_align)

            for t in doc.tables:
                t.alignment = WD_TABLE_ALIGNMENT.RIGHT
                tblPr = t._tbl.tblPr
                for el in tblPr.findall(qn('w:bidiVisual')):
                    tblPr.remove(el)
                tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))

                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        is_func_domain_cell = (len(row.cells) == 3 and c_idx == 0 and r_idx > 0)
                        for p in cell.paragraphs:
                            is_p_centered = (p.alignment == WD_ALIGN_PARAGRAPH.CENTER or r_idx == 0)
                            for r in p.runs:
                                if is_func_domain_cell:
                                    apply_rPr_david_rtl(r, bold=True)
                                else:
                                    apply_rPr_david_rtl(r)
                            apply_pPr_rtl(p, align="center" if is_p_centered else "right")

            bio = io.BytesIO()
            doc.save(bio)
            return bio.getvalue()

        # הפקה מאפס במידה ואין קובץ טמפלייט
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.65)
            section.right_margin = Inches(0.65)
            sectPr = section._sectPr
            sectPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))

        p_t = doc.add_paragraph()
        r = p_t.add_run("תוכנית לימודים אישית לתלמיד בחינוך המיוחד – תל\"א")
        apply_rPr_david_rtl(r, bold=True, size_pt=16.0)
        apply_pPr_rtl(p_t, align="center", space_after_pt=12)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    st.markdown('<span class="marker-download"></span>', unsafe_allow_html=True)
    st.markdown('<div class="download-btn">', unsafe_allow_html=True)
    st.download_button(
        label="הורד קובץ Word מעוצב עם טבלאות מלאות",
        data=build_template_docx(),
        file_name=f"תלא_{active_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.markdown('</div>', unsafe_allow_html=True)
